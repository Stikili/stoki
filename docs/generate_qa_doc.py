"""
Generate the Stoki QA Test Plan as a Word document.

Uses python-docx to produce a properly-formatted .docx file that a QA
engineer can open, edit, and hand back with results marked. Structure
mirrors the app's actual feature surface as of 2026-07-14.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r"c:\Users\sbuti\source\repos\stoki\docs\Stoki-QA-Test-Plan.docx"

# ── Brand tokens ────────────────────────────────────────────────────────
BRAND_GREEN = RGBColor(0x00, 0xC8, 0x96)
DEEP_INK    = RGBColor(0x0A, 0x0E, 0x17)
MID_GREY    = RGBColor(0x4A, 0x58, 0x78)
LIGHT_GREY  = RGBColor(0x7B, 0x8C, 0xA1)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BRAND_GREEN
        r.font.size = Pt(20)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = DEEP_INK
        r.font.size = Pt(15)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = DEEP_INK
        r.font.size = Pt(12)
    return p

def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def checkbox(text):
    """A test-case bullet with a leading checkbox."""
    return doc.add_paragraph(f'☐  {text}', style='List Bullet')

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    header_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = ''
        run = header_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════

# Empty space at top
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('STOKI')
r.bold = True
r.font.size = Pt(48)
r.font.color.rgb = BRAND_GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('QA Test Plan')
r.font.size = Pt(24)
r.font.color.rgb = DEEP_INK

doc.add_paragraph()
doc.add_paragraph()

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run('Comprehensive testing scope for the Stoki business assistant.')
r.font.size = Pt(12)
r.font.color.rgb = MID_GREY
r.italic = True

for _ in range(6):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Version 1.0  ·  July 2026\nStoki (Pty) Ltd  ·  Cape Town, South Africa\n\nProduction: https://www.stokiapp.com\nSupport: support@stokiapp.com')
r.font.size = Pt(11)
r.font.color.rgb = LIGHT_GREY

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. ABOUT STOKI
# ═══════════════════════════════════════════════════════════════════════

h1('1. About Stoki')

para(
    'Stoki is an AI-powered business assistant built for South African shops '
    'and SMMEs. It combines a full retail till, credit book, inventory, '
    'invoicing, payables, payroll, fixed-asset register, cashflow forecasting '
    'and SARS-ready tax reporting with an AI advisor that is grounded in the '
    'user\'s data and the wider SA economy. The product is accessible on the '
    'web (PWA, works offline) and via WhatsApp.'
)

h3('1.1 Target audience')
bullet('Informal traders — spaza shops, general dealers, food stalls')
bullet('Service businesses — salons, plumbers, mobile operators, small contractors')
bullet('Growing SMMEs — VAT-registered, 1-10 employees, formal operations')

h3('1.2 Tech stack (context for reproducing environments)')
bullet('Framework: Next.js 16 (App Router) + React 19 + TypeScript')
bullet('Database + Auth: Supabase (Postgres with RLS)')
bullet('AI: Anthropic Claude via @anthropic-ai/sdk with tool-use')
bullet('Email: Resend SMTP (transactional) + Zoho Mail (support inbox)')
bullet('Errors: Sentry (@sentry/nextjs, DE region)')
bullet('Hosting: Vercel Pro + Vercel DNS on stokiapp.com')
bullet('Testing: Vitest 4 (unit) + Playwright 1.61 (E2E)')

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. TEST ENVIRONMENTS
# ═══════════════════════════════════════════════════════════════════════

h1('2. Test Environments')

table(
    ['Environment', 'URL', 'Purpose', 'Data'],
    [
        ['Production', 'https://www.stokiapp.com', 'Live app used by real users', 'Real customer data — do not seed test data'],
        ['Preview', 'https://stoki-app-<hash>.vercel.app', 'Per-PR deploys from Vercel', 'Points at production Supabase — test carefully'],
        ['Local', 'http://localhost:3000', 'Developer machine + your own laptop', 'Points at your Supabase project only'],
    ],
    col_widths=[3.0, 5.0, 4.0, 5.0],
)

h3('2.1 Getting test accounts')

para(
    'The founder (sbutikili@gmail.com) is currently the sole /admin. To be '
    'issued a test account:'
)
bullet('Ask the founder to visit https://www.stokiapp.com/admin')
bullet('Enter your email in the "Beta invite" form and click Create')
bullet('Founder shares the auto-generated temp password with you over WhatsApp')
bullet('You log in and change the password in Settings → Account & preferences')

h3('2.2 Supported browsers')

bullet('Chrome (Windows / macOS / Android)  — primary')
bullet('Safari (macOS / iOS)                  — primary')
bullet('Edge (Windows)                       — full support')
bullet('Firefox (Windows / macOS)            — best-effort')
bullet('Samsung Internet (Android)           — best-effort')

h3('2.3 Supported viewports')

table(
    ['Class', 'Width', 'Notes'],
    [
        ['iPhone SE',       '375px', 'Smallest supported — no vertical scroll in onboarding'],
        ['Mobile',          '390-430px', 'iPhone 14-15 Pro / Pixel'],
        ['Tablet portrait', '768px', 'iPad'],
        ['Tablet landscape', '1024px', 'iPad landscape'],
        ['Desktop',         '1280-1440px', 'MacBook, standard laptop'],
        ['Wide',            '1920px+', 'External monitor'],
    ],
    col_widths=[4.0, 3.0, 8.0],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. TEST SCOPE
# ═══════════════════════════════════════════════════════════════════════

h1('3. Test Scope by Feature Area')

para(
    'Each subsection lists concrete test cases in the form of a checklist. '
    'Tick the box when confirmed. Add a bug reference (e.g., "BUG-042") '
    'inline for any failed case.',
    italic=True,
)

# ── 3.1 Landing page ───────────────────────────────────────────────────

h2('3.1 Landing page (public)')

h3('URL: /  (unauthenticated)')

checkbox('Page loads within 3 seconds on desktop and mobile')
checkbox('Hero renders: overline, headline in white + emerald accent, sub-copy')
checkbox('"Sign up free" button visible on all three slides')
checkbox('"Sign in" button in top-right of header')
checkbox('Theme toggle switches between dark and light mode; preference persists after refresh')
checkbox('Mobile: three slides are horizontally swipeable with snap points; pagination dots update')
checkbox('Desktop: same three slides stack vertically with normal page scroll')
checkbox('Features grid shows 6 cards (Money in & out, SARS-ready, Cash flow, Stoki AI, POS, Stock)')
checkbox('Proof strip shows 4 chips (5 SA banks, PAYE·UIF·SDL, 30-day forecast, Works offline)')
checkbox('CTA slide: hello@stokiapp.com and support@stokiapp.com mailto links open the mail client')
checkbox('Company registration text visible: "Stoki (Pty) Ltd · Reg. K2026258855"')
checkbox('Privacy Policy and Terms links open respective pages')

# ── 3.2 Authentication ─────────────────────────────────────────────────

h2('3.2 Authentication')

h3('3.2.1 Sign up (email + password)')

checkbox('URL /login?intent=register opens in "Create your account" mode')
checkbox('Fields: Email, Password, Confirm password, consent checkbox')
checkbox('Password field is masked by default; eye icon toggles visibility')
checkbox('"Create account" button disabled until email + password (8+ chars) + confirm match + consent are all valid')
checkbox('Submit with valid data → success toast: "Account created — check your email to confirm"')
checkbox('Confirmation email arrives at the supplied address within 30 seconds')
checkbox('Sender is noreply@stokiapp.com (via Resend)')
checkbox('Clicking the confirmation link lands the user on /dashboard signed in')
checkbox('Attempting to register an already-registered email surfaces a clear error')

h3('3.2.2 Sign in (email + password)')

checkbox('URL /login opens in "Welcome back" mode by default')
checkbox('"Remember me" checkbox pre-fills email on next visit if checked')
checkbox('"Forgot password?" opens the reset flow')
checkbox('Valid credentials redirect to /dashboard')
checkbox('Invalid credentials show "Wrong email or password" (does not leak which is wrong)')
checkbox('Session persists across browser tabs')
checkbox('Session survives a browser restart (if "Remember me" was checked)')

h3('3.2.3 Password reset')

checkbox('"Forgot password?" opens a full-page reset form')
checkbox('Submitting an email sends a reset link within 30 seconds')
checkbox('Reset link opens /auth/reset-password with a valid session')
checkbox('New password must be 8+ chars')
checkbox('After reset, user is redirected to /dashboard signed in')

h3('3.2.4 Phone OTP (currently hidden)')

para(
    'The phone OTP path is intentionally hidden until Twilio is wired in Supabase. '
    'Verify it is NOT reachable:',
    italic=True,
)
checkbox('Login page does not show an "OTP — SA cell number" alt-button')
checkbox('No divider is rendered below the primary CTA')
checkbox('Mode is always "email"; a saved phone identifier does not switch modes')

# ── 3.3 Onboarding ─────────────────────────────────────────────────────

h2('3.3 Onboarding (post-signup)')

h3('URL: /onboarding — 7-panel swipe deck')

para(
    'Panels: Type → Name → Location → Cash → Business → Usage → Pack. '
    'Test on both mobile (swipe) and desktop (click).',
)

h3('3.3.1 Panel 1 — Business type')

checkbox('4 categories visible: Spaza shop, General dealer, Food stall, Other')
checkbox('Selecting a category shows an emerald tick')
checkbox('"Continue" button advances to Panel 2')
checkbox('Swipe-left on mobile also advances')
checkbox('Back button is disabled on this first panel')

h3('3.3.2 Panel 2 — Store name + phone')

checkbox('Store name field is required; Continue is disabled if empty')
checkbox('Placeholder adapts to the chosen category ("Thabo\'s Spaza" vs "Mama\'s Kitchen")')
checkbox('Phone field is optional')
checkbox('Continue creates the store and advances (spinner shown during save)')
checkbox('Swipe-left also triggers the save action (validation applied)')
checkbox('Back button returns to Panel 1 without losing the typed name/phone')

h3('3.3.3 Panel 3 — Location (GPS)')

checkbox('"Use my location" button prompts for browser geolocation permission')
checkbox('Latitude and Longitude fields populate on permission grant')
checkbox('Manual lat/lng entry works')
checkbox('Denying permission shows a friendly error, allows manual entry')
checkbox('"Swipe to skip" text is visible; skipping does not fail the flow')

h3('3.3.4 Panel 4 — Cash on hand')

checkbox('Number input, R0 accepted, empty accepted')
checkbox('Swipe to skip works')

h3('3.3.5 Panel 5 — About your business')

checkbox('Two toggles: "I\'m VAT-registered" and "I have employees on payroll"')
checkbox('Both default to off')
checkbox('Toggling on and off updates state immediately')

h3('3.3.6 Panel 6 — Dashboard density')

checkbox('Three options: "Run my till and stock", "Manage the business books too", "I\'m not sure yet"')
checkbox('Radio-style selection; only one active')
checkbox('Default is "I\'m not sure yet"')

h3('3.3.7 Panel 7 — Starter pack')

checkbox('Card previews products included for the chosen category')
checkbox('"Yes, load starter pack" shows for non-"Other" categories')
checkbox('"I\'ll add my own" always available')
checkbox('For "Other" category, only "Let\'s go →" shows (no starter pack option)')
checkbox('Choosing either option completes onboarding and lands on /dashboard')

h3('3.3.8 Overall onboarding UX')

checkbox('No vertical scroll on any panel at iPhone SE (375×667) viewport')
checkbox('Desktop: card is centered horizontally, max-width ~448px, max-height ~820px')
checkbox('Progress dots at the top update as you advance')
checkbox('Wordmark visible above the progress dots')
checkbox('Back button greys out on first panel; Continue disabled when validation fails')
checkbox('Swipe gestures do not trigger while typing in an input')

page_break()

# ── 3.4 Dashboard ──────────────────────────────────────────────────────

h2('3.4 Dashboard')

h3('URL: /dashboard')

h3('3.4.1 Layout')

checkbox('Greeting adapts to time of day (Good morning / afternoon / evening)')
checkbox('Store name displayed under greeting')
checkbox('Trial banner appears only in the final 7 days of the 90-day trial')
checkbox('Setup checklist appears while any item is undone; disappears once complete')
checkbox('Revenue hero shows today\'s revenue in ZAR with 2 decimals')
checkbox('Revenue comparison ("vs same day last week") displays when historical data exists')
checkbox('Status pills row: alerts, low stock, expiring, credit owed (each hides when count = 0)')
checkbox('7-day chart renders bars; today\'s bar is highlighted in emerald')
checkbox('Ask Stoki card removed (assistant only accessible via BottomNav center button)')
checkbox('Manage tile grid split into "Daily" and "Books" sections')
checkbox('Books section is collapsed by default when simple_view = true (default)')
checkbox('Books section shows "Show N ⌄" toggle; clicking expands the tiles')
checkbox('Bottom nav has 5 items: Home, Credit, Stoki FAB (center), Sales, Inventory')
checkbox('Stoki FAB is raised, emerald, larger than side items, links to /advisor')

h3('3.4.2 Manage tiles')

para('The following 16 tiles are expected. Confirm each shows up per section:')

table(
    ['Tile', 'Section', 'Plan required'],
    [
        ['Cash up',   'Daily', 'Free'],
        ['Expenses',  'Daily', 'Free'],
        ['Suppliers', 'Daily', 'Free'],
        ['Stocktake', 'Daily', 'Free'],
        ['Airtime',   'Daily', 'Free'],
        ['Prices',    'Daily', 'Free'],
        ['Reports',   'Books', 'Free'],
        ['Invoices',  'Books', 'Pro'],
        ['Customers', 'Books', 'Pro'],
        ['Payables',  'Books', 'Pro'],
        ['Cash flow', 'Books', 'Free'],
        ['Assets',    'Books', 'Business'],
        ['POs',       'Books', 'Pro'],
        ['Payroll',   'Books', 'Business'],
        ['Reconcile', 'Books', 'Pro'],
        ['Broadcasts','Books', 'Business'],
    ],
    col_widths=[5.0, 3.5, 4.0],
)

checkbox('Each tile shows an ⓘ badge in the top-left corner')
checkbox('Tapping ⓘ opens a hint modal with one-sentence description; does NOT navigate')
checkbox('Long-press (450ms) on the tile body also opens the hint modal')
checkbox('Cashier role sees only "Prices" tile from this grid')
checkbox('Tiles above the user\'s current plan show a lock badge in top-right')
checkbox('Clicking a locked tile opens the upgrade prompt (not the tile\'s page)')
checkbox('When simpleView = true, Books section is collapsed on first render')
checkbox('When simpleView = false, both sections are expanded on first render')

# ── 3.5 Sales / till ──────────────────────────────────────────────────

h2('3.5 Sales / till')

h3('URL: /sales — plus the FAB from any page')

checkbox('FAB (green centre button in BottomNav) opens the sale entry sheet')
checkbox('Product search field filters by name substring')
checkbox('Selecting a product adds it to the cart')
checkbox('Quantity can be adjusted with +/- controls')
checkbox('Cart total updates in real time')
checkbox('Weighable items prompt for a mass entry (kg to 3 decimals)')
checkbox('Bundle items expand into components at cost/price allocation')
checkbox('Payment method selector: Cash, Card, EFT, Credit')
checkbox('Cash payment prompts for tendered amount and calculates change')
checkbox('Airtime PIN dispenses automatically for airtime SKUs; PIN visible once')
checkbox('Completing a sale returns to the previous page with a success toast')
checkbox('Recorded sale appears in the day\'s revenue hero within 30 seconds')
checkbox('Returns / refunds: negative-quantity or dedicated return flow')

h3('3.5.1 Cash up (end of day)')

checkbox('Dashboard → Manage → Daily → Cash up tile opens the flow')
checkbox('Expected cash calculated from today\'s cash sales')
checkbox('User enters actual cash count')
checkbox('Variance is displayed (over / short / matched)')
checkbox('Saving locks the day and creates a cashup record')

# ── 3.6 Inventory ──────────────────────────────────────────────────────

h2('3.6 Inventory')

h3('URL: /inventory')

checkbox('Product list loads within 2 seconds')
checkbox('Search filters by name substring')
checkbox('Filter for "low stock", "out of stock", "expiring" work independently')
checkbox('Add product: name, price, cost, quantity, reorder point, expiry date, VAT flag')
checkbox('Edit product: same fields, changes persist')
checkbox('Delete product: soft-delete with confirmation prompt')
checkbox('Barcode scan (camera) via html5-qrcode; permission prompt shown first')
checkbox('Restock: increases qty, records the transaction, updates last-restocked-at')
checkbox('Wastage: decreases qty with a reason (expired / damaged / stolen)')
checkbox('Bundle configuration: add components with per-unit qty and cost split')
checkbox('Bulk import via starter pack works and dedupes on name match')
checkbox('Low-stock alert triggers when qty <= reorder_point')
checkbox('Expiry alert triggers when days-until-expiry <= 7')

h3('3.6.1 Stocktake')

checkbox('Stocktake flow lists all active products')
checkbox('Enter counted qty per product')
checkbox('Variance shown per product (system vs counted)')
checkbox('Save adjusts system qty to counted qty and logs a variance entry')

# ── 3.7 Credit book ───────────────────────────────────────────────────

h2('3.7 Credit book / debtors')

h3('URL: /credit')

checkbox('Debtors list ordered by amount owed (largest first)')
checkbox('"Total outstanding" summary at top matches sum of individual debts')
checkbox('Add debtor: name, phone, opening balance')
checkbox('Record credit transaction: increases balance, logs entry with date')
checkbox('Record payment: decreases balance, logs entry')
checkbox('Aging: 0-30, 31-60, 61-90, 90+ day buckets calculated correctly')
checkbox('Overdue debtors flagged with red pill')
checkbox('Debtor detail page shows full transaction history')
checkbox('WhatsApp reminder button opens a pre-filled WhatsApp message with the amount owed')

# ── 3.8 Suppliers / payables / POs ────────────────────────────────────

h2('3.8 Suppliers, Payables, Purchase Orders (Pro)')

h3('URL: /suppliers · /payables · /purchase-orders')

checkbox('Suppliers list shows all recorded suppliers')
checkbox('Supplier detail shows last-90-days restock history')
checkbox('Payables aging matches invoice aging buckets')
checkbox('PO creation: supplier, line items, expected delivery date')
checkbox('PO status transitions: draft → sent → partially received → received')
checkbox('Goods received against PO decrements outstanding qty')
checkbox('Free-tier users see lock badges on these tiles + upgrade prompt on click')

# ── 3.9 Invoices / customers ──────────────────────────────────────────

h2('3.9 Invoices & B2B customers (Pro)')

h3('URL: /invoices · /customers')

checkbox('Add customer: name, contact, email, phone, VAT number, billing address, payment terms')
checkbox('Customer VAT number validation (SA format: 10 digits, starts with 4)')
checkbox('Create invoice: pick customer, line items, VAT calculated per line')
checkbox('Invoice numbers sequential per store')
checkbox('SARS tax-invoice fields present: seller VAT, buyer VAT, invoice date, invoice number, VAT breakdown')
checkbox('Invoice status: draft → sent → partially paid → paid')
checkbox('Send invoice via WhatsApp opens a pre-filled message with invoice PDF link')
checkbox('Send invoice via email opens the user\'s mail client with subject/body pre-filled')
checkbox('Record payment against invoice; overpayment blocked')
checkbox('Overdue invoices display red pill after payment_terms + 1 day')

# ── 3.10 Reports ──────────────────────────────────────────────────────

h2('3.10 Reports')

h3('URL: /reports')

checkbox('P&L (income statement) for a given period')
checkbox('VAT201 worksheet with block 1 (standard-rated), block 2 (zero-rated), etc.')
checkbox('Sales detail report with per-product breakdown')
checkbox('Provisional tax estimate matches SA tax tables for the taxpayer type')
checkbox('CSV export downloads a file with correct headers and row counts')
checkbox('PDF export renders correctly and is printable')
checkbox('Cross-store report (Business only) aggregates multiple stores')
checkbox('Cross-store report locked for Pro users; shows upgrade prompt')

# ── 3.11 Bank reconciliation ──────────────────────────────────────────

h2('3.11 Bank reconciliation (Pro)')

h3('URL: /reconcile')

checkbox('CSV upload accepts files from all 5 SA banks: FNB, Standard, ABSA, Nedbank, Capitec')
checkbox('Header row detection works across bank formats')
checkbox('Auto-match against existing invoices and expenses on amount + date proximity')
checkbox('Manual match: user drags/clicks a bank transaction onto an invoice/expense')
checkbox('Unmatched transactions can be created as new expenses in one click')
checkbox('Reconciliation summary: total matched vs unmatched')

# ── 3.12 Payroll ──────────────────────────────────────────────────────

h2('3.12 Payroll (Business)')

h3('URL: /payroll')

checkbox('Add employee: name, ID number (hashed), start date, salary, PAYE bracket, UIF applicable, SDL applicable')
checkbox('Employee ID number stored as hash; only last 4 digits visible')
checkbox('Run monthly payroll: PAYE, UIF, SDL calculated per SARS 2026 tables')
checkbox('Employer UIF and SDL contributions calculated correctly')
checkbox('EMP201 XML export downloads')
checkbox('Payslip PDF generated per employee, with tax breakdown')
checkbox('Payroll run logs to journal and appears in P&L as an expense line')
checkbox('Payroll tile locked for Free/Pro users; visible with lock badge')

# ── 3.13 Fixed assets ─────────────────────────────────────────────────

h2('3.13 Fixed assets (Business)')

h3('URL: /assets')

checkbox('Add asset: name, category, purchase date, cost, useful life (years)')
checkbox('Depreciation method: straight-line (default) with monthly posting')
checkbox('Asset register shows carrying value, accumulated depreciation, remaining life')
checkbox('Monthly cron posts depreciation to P&L automatically')
checkbox('Asset disposal: sale price, book value, gain/loss calculated')
checkbox('Locked for Free/Pro users')

# ── 3.14 Cash flow forecast ───────────────────────────────────────────

h2('3.14 Cash flow forecast')

h3('URL: /cashflow')

checkbox('30-day forecast chart renders')
checkbox('Assumptions displayed (starting cash, avg daily revenue, avg daily variable expense)')
checkbox('Open invoices contribute inflows on their due date')
checkbox('Open supplier bills contribute outflows on their due date')
checkbox('Recurring expenses (rent, electricity) contribute outflows on their next posting date')
checkbox('Deficit day flagged in red on the chart')
checkbox('Deficit banner appears on the dashboard when a 14-day forecast shows a shortfall')

# ── 3.15 Expenses ─────────────────────────────────────────────────────

h2('3.15 Expenses')

h3('URL: /expenses')

checkbox('Add expense: date, amount, category, supplier (optional), reference')
checkbox('Recurring expense setup: rent, electricity, insurance with monthly/weekly cadence')
checkbox('Cron auto-posts recurring expenses on their schedule (verify by advancing system date)')
checkbox('Expense category totals in reports')

# ── 3.16 WhatsApp bot ────────────────────────────────────────────────

h2('3.16 WhatsApp bot (Meta Cloud API)')

para(
    'Bot is accessed by sending a message to the Stoki WhatsApp number '
    'linked to the user\'s store. First message from a linked number is '
    'routed to the user\'s store context automatically.',
    italic=True,
)

h3('3.16.1 Help fast-path (no LLM call)')

checkbox('Sending "help", "menu", "hi", "hey", "hello", "start", "?", "sawubona", or "molo" returns the deterministic help menu within 3 seconds')
checkbox('Help menu includes example queries (business, profit, reorder, debtors, sell, expense, economy, how-to)')
checkbox('Help menu mentions support@stokiapp.com as the human escalation path')

h3('3.16.2 Data queries (LLM + tool use)')

checkbox('"How is business today?" — returns today\'s revenue, txn count, low stock summary')
checkbox('"What\'s my profit this month?" — returns month-to-date margin')
checkbox('"What should I reorder?" — surfaces low-stock items ordered by turnover')
checkbox('"Who owes me money?" — lists debtors with outstanding balances')

h3('3.16.3 Actions (LLM + tool use)')

checkbox('"Sold 5 bread" — fuzzy-matches product, records the sale, confirms in the reply')
checkbox('"Paid 200 rand for airtime" — creates an expense in the airtime category')
checkbox('"I restocked 2 cases of Coke" — increases stock and logs a restock entry')
checkbox('"Throw away expired bread" — creates a wastage entry')
checkbox('Unrecognised product name → bot lists close matches for user to confirm')

h3('3.16.4 Economy queries (web search)')

checkbox('"How is the economy?" — pulls market context + web sources; cites URLs')
checkbox('"Is fuel increasing?" — searches aa.co.za or energy.gov.za; returns date + rand-per-litre move')
checkbox('"What did SARB say today?" — searches resbank.co.za and cites the specific announcement')

h3('3.16.5 How-to queries')

checkbox('"How do I do cash-up?" — bot answers with in-app nav (Dashboard → Manage → Daily → Cash up)')
checkbox('"Where do I find VAT?" — bot points at Dashboard → Manage → Books → Reports/VAT')
checkbox('"How do I invite a teammate?" — bot points at Settings → Team → Invite member (owner only)')

h3('3.16.6 Voice + multilingual')

checkbox('Voice notes are transcribed and processed same as text messages')
checkbox('isiZulu message ("Ngithengise 5 isinkwa") is understood as "sold 5 bread"')
checkbox('Afrikaans message ("Verkoop 5 brood") is understood as "sold 5 bread"')

h3('3.16.7 Rate limiting')

checkbox('IP throttle: WhatsApp webhook accepts up to 120 requests/min per IP; 429 above')
checkbox('Global daily budget: 50,000 advisor calls per day; blocks further calls once exceeded')
checkbox('Per-user daily limit: 20 messages/day on Free, 500 on Pro, 5000 on Business')

# ── 3.17 AI advisor (in-app) ─────────────────────────────────────────

h2('3.17 AI advisor (in-app)')

h3('URL: /advisor')

checkbox('Chat interface loads with suggested prompts')
checkbox('Send a message → typing indicator appears → response streams back')
checkbox('Conversation memory: previous turns visible in scroll')
checkbox('Advanced insight questions gate on plan; Free users see lock messages for the 7 Pro-only insights')
checkbox('Rate limit: 30 messages/minute per user (Pro cap)')

# ── 3.18 Broadcasts ──────────────────────────────────────────────────

h2('3.18 WhatsApp broadcasts (Business)')

h3('URL: /broadcasts')

checkbox('Template creation form: template name, category, body, media (optional)')
checkbox('Meta-template approval status displayed (pending / approved / rejected)')
checkbox('Recipient selection: opted-in customer segments')
checkbox('Send broadcast: rate-limited, batches sent in the background')
checkbox('Delivery report: sent / delivered / read / failed counts')
checkbox('Locked for Free/Pro users')

# ── 3.19 Settings ────────────────────────────────────────────────────

h2('3.19 Settings')

h3('URL: /settings')

checkbox('Settings index shows sections filtered by role (cashier sees only Account & preferences)')
checkbox('Support card visible: two side-by-side buttons for Ask Stoki (advisor) + Email us (support@stokiapp.com)')
checkbox('Version and region shown at bottom')
checkbox('Download data backup (JSON) downloads a file with the store\'s data')

h3('3.19.1 Store details')

checkbox('Edit name, phone, area/suburb, WhatsApp number, business address')
checkbox('GPS capture button works; manual lat/lng entry works')
checkbox('Cash on hand editable; last-updated stamp shown')
checkbox('Delete store button appears only if user has 2+ stores; confirms before deleting')

h3('3.19.2 Team')

checkbox('Owner sees the Team section; cashier does not')
checkbox('Invite by email: existing user is added instantly; new user gets a magic-link invite')
checkbox('Role selector: owner / manager / cashier')
checkbox('Role change from row-level menu works')
checkbox('Remove member works; owner cannot remove themselves')
checkbox('Plan gate: Pro allows 1 teammate, Business allows up to 5, Enterprise unlimited')

h3('3.19.3 VAT')

checkbox('Toggle VAT registration on/off')
checkbox('VAT number field appears when registered')
checkbox('VAT rate editable (default 15%)')
checkbox('Taxpayer type selector: sole_prop / sbc / turnover_tax / company')
checkbox('Bulk-set all products to VAT-inclusive / VAT-exclusive works')

h3('3.19.4 Account & preferences')

checkbox('Add / change email (sends confirmation)')
checkbox('Theme toggle (dark / light) applied immediately; preference persists')
checkbox('Language selector: English, isiZulu, isiXhosa, Afrikaans (whatever is wired)')
checkbox('Push notifications: subscribe / unsubscribe works; test notification arrives')
checkbox('Dashboard density: Simple / Full toggle; setting persists per store')
checkbox('Sign out button ends session and redirects to /login')
checkbox('Delete my account: two confirmation prompts + "type DELETE" step; account is deleted')

h3('3.19.5 Billing & plan')

checkbox('Current plan shown correctly (Free / Pro / Business / Enterprise)')
checkbox('Trial pill and days remaining shown when trial is active')
checkbox('Four plan cards visible: Free, Pro (R99/mo), Business (R249/mo), Enterprise (from R899/mo)')
checkbox('Highlighted card shows the current plan')
checkbox('CTA buttons (until Ozow is wired) open pre-filled mailto to support@stokiapp.com')
checkbox('Custom-plan email link at bottom points to support@stokiapp.com')

h3('3.19.6 Recurring expenses')

checkbox('Add rule: name, amount, category, cadence (weekly/monthly), start date')
checkbox('Active rules listed with next-post date')
checkbox('Deactivate rule pauses future postings')

h3('3.19.7 Market values')

checkbox('Owner only')
checkbox('Displays SARB rate, fuel prices, CPI, USD/ZAR (read-only from source cache)')

# ── 3.20 Multi-store ─────────────────────────────────────────────────

h2('3.20 Multi-store (Business)')

checkbox('Store header shows current store name; click reveals a switcher')
checkbox('Switching stores re-renders dashboard with the new store\'s data')
checkbox('/stores lists all stores owned by the user')
checkbox('"Add store" is disabled for Free/Pro users (locked, upgrade prompt)')
checkbox('Business tier allows up to 3 stores; 4th triggers upgrade prompt to Enterprise')
checkbox('Cross-store queries (reports) show data across all stores')
checkbox('RLS enforced: user cannot query data from a store they don\'t belong to')

# ── 3.21 Trial + subscription ────────────────────────────────────────

h2('3.21 Trial & subscription lifecycle')

checkbox('New signup receives 90-day Business-tier access automatically')
checkbox('effectivePlan() returns "business" for trial users on Free plan')
checkbox('effectivePlan() does not downgrade existing Business/Enterprise users during trial')
checkbox('Trial banner appears on dashboard at day 83 (7 days remaining)')
checkbox('Trial banner urgency style flips at day 87 (3 days remaining) — amber to emerald')
checkbox('Trial expiry: user drops to Free automatically; no data loss')
checkbox('After expiry, previously visible paid tiles show lock badges')
checkbox('CTAs on locked tiles open the upgrade prompt (mailto until Ozow lands)')

# ── 3.22 Plan gating ────────────────────────────────────────────────

h2('3.22 Plan gating (enforcement)')

para(
    'These tests confirm the plan-gate boundaries are enforced both in UI '
    '(lock badges) and server actions (returned errors when a locked feature '
    'is invoked directly).',
)

table(
    ['Feature', 'Gate ID', 'Min plan'],
    [
        ['B2B invoicing',       'invoice.create',           'Pro'],
        ['Bank reconciliation', 'reports.bank_reconcile',   'Pro'],
        ['Supplier payables',   'payables.manage',          'Pro'],
        ['Purchase orders',     'purchase_orders.create',   'Pro'],
        ['Payroll',             'payroll.run',              'Business'],
        ['Fixed assets',        'assets.manage',            'Business'],
        ['WhatsApp broadcasts', 'broadcast.send',           'Business'],
        ['Team beyond owner',   'team.invite.beyond_owner', 'Pro'],
        ['Team beyond 2',       'team.invite.beyond_2',     'Business'],
        ['Multi-store',         'store.create.beyond_1',    'Business'],
        ['Unlimited stores',    'store.create.beyond_3',    'Enterprise'],
        ['Cross-store reports', 'reports.cross_store',      'Business'],
    ],
    col_widths=[6.0, 5.5, 3.0],
)

checkbox('For each gate above, verify the tile/entry point shows a lock badge for a plan below the min')
checkbox('Click on locked tile opens UpgradePrompt with the correct label + description')
checkbox('Direct API/server-action call to the locked feature returns a FeatureLockedError')

# ── 3.23 Admin console ─────────────────────────────────────────────────

h2('3.23 Admin console (/admin)')

para(
    'Gated by ADMIN_EMAILS env var. Only whitelisted emails can reach this '
    'page.',
    italic=True,
)

checkbox('Non-admin logged-in user gets 404 for /admin')
checkbox('Unauthenticated visitor is redirected to /login')
checkbox('Admin sees Aggregate stats: total signups, active stores, on-trial, paid')
checkbox('Beta invite card at top: enter email → Create → temp password + WhatsApp message rendered once')
checkbox('Temp password copy-to-clipboard works')
checkbox('WhatsApp message copy-to-clipboard works')
checkbox('Duplicate email is rejected with a clear error')
checkbox('Table lists all signups with email, phone, store, category, plan, onboarding status, signed-up-when')
checkbox('Sort/filter behavior (if applicable)')

# ── 3.24 PWA / offline ─────────────────────────────────────────────────

h2('3.24 PWA & offline')

checkbox('Install prompt appears on Chrome/Edge mobile visit; iOS shows Share → Add to Home Screen guidance')
checkbox('Installing the PWA creates an app icon on the home screen')
checkbox('Opening the PWA launches in standalone mode (no browser chrome)')
checkbox('Airplane mode / offline: dashboard still renders from cache')
checkbox('Offline banner appears when connection drops')
checkbox('Actions taken offline are queued and synced when connection returns')
checkbox('Push notification arrives when a critical alert fires')

# ── 3.25 Compliance / security ────────────────────────────────────────

h2('3.25 Compliance & security')

checkbox('All URLs served over HTTPS with a valid certificate')
checkbox('Session cookies are httpOnly and Secure')
checkbox('Cross-tenant test: try to fetch another store\'s product by ID via the browser network tab; RLS blocks it')
checkbox('Rate limits enforced: rapid-fire advisor requests return 429 with Retry-After')
checkbox('SQL injection: attempt to inject SQL via free-text inputs (product names, invoice descriptions); nothing executes')
checkbox('XSS: attempt to inject <script> in a product name; renders as text, not code')
checkbox('CSRF: server actions require valid Next.js session cookie')
checkbox('Employee ID number is stored hashed; only last 4 visible in UI')
checkbox('Service role Supabase key never appears in browser network responses')
checkbox('Privacy policy: link works, mentions POPIA, lists Information Officer contact (support@stokiapp.com)')
checkbox('Terms of service: link works, dispute resolution contact = support@stokiapp.com')

# ── 3.26 Performance ─────────────────────────────────────────────────

h2('3.26 Performance benchmarks')

table(
    ['Metric', 'Target', 'Notes'],
    [
        ['Time to First Byte',          '< 800ms',  'Vercel edge; measure via DevTools Network'],
        ['Largest Contentful Paint',    '< 2.5s',   'Dashboard first paint'],
        ['Time to Interactive',         '< 3.5s',   'Dashboard fully clickable'],
        ['Bundle size (compressed)',    '< 250KB',  'Landing page first load'],
        ['Advisor response latency',    '< 8s',     'p95 for standard queries'],
    ],
    col_widths=[5.5, 3.0, 6.5],
)

checkbox('Run Lighthouse on /, /dashboard, /login; screenshot the results')
checkbox('Run Lighthouse on mobile emulation (Slow 4G, mid-tier phone)')
checkbox('Test on a real budget SA phone (Nokia G20 / Samsung A03 class device)')

# ── 3.27 Error monitoring ────────────────────────────────────────────

h2('3.27 Error monitoring (Sentry)')

checkbox('Trigger a client-side error (browser console: setTimeout(() => { throw new Error("qa test") }, 0))')
checkbox('Error appears in Sentry Issues within 60 seconds')
checkbox('Error includes environment tag "production" and release SHA')
checkbox('Trigger a server-side error (best via a deliberately-crashing route in preview)')
checkbox('Server error is captured by onRequestError hook with route context')
checkbox('Global-error page renders when React tree crashes; support mailto works')

# ── 3.28 Email delivery ──────────────────────────────────────────────

h2('3.28 Email delivery (Resend + Zoho)')

checkbox('Sign-up confirmation email arrives within 30 seconds')
checkbox('Password reset email arrives within 30 seconds')
checkbox('Emails sent from noreply@stokiapp.com; not marked as spam by Gmail/Outlook')
checkbox('Support inbox: sending mail to support@stokiapp.com from external inbox lands in Zoho web')
checkbox('Support inbox: push notification arrives on the phone (Zoho Mail app)')
checkbox('Welcome inbox: sending mail to hello@stokiapp.com lands in the same Zoho inbox')

# ── 3.29 Cross-cutting UX ────────────────────────────────────────────

h2('3.29 Cross-cutting UX')

checkbox('Empty states: every list page (Products, Debtors, Invoices, etc.) shows a helpful empty-state card')
checkbox('Loading states: skeleton or spinner during data fetch')
checkbox('Error states: friendly error message + retry option when a request fails')
checkbox('Form validation: real-time feedback on required fields, format errors')
checkbox('Focus management: modals return focus to the trigger on close')
checkbox('Keyboard navigation: Tab order is logical across all pages')
checkbox('Screen reader: main landmarks (nav, main, footer) present')
checkbox('Dark mode: no white flash on page load')
checkbox('Long product names truncate with ellipsis; hover shows full name')

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. TESTING METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════

h1('4. Testing Methodology')

h2('4.1 Test types')

table(
    ['Type', 'When', 'Owner'],
    [
        ['Smoke',       'Every deploy',                     'Automated (Playwright) + QA sanity check'],
        ['Regression',  'Weekly / major release',           'QA — full test plan'],
        ['Exploratory', 'Weekly',                           'QA — creative attempts to break the app'],
        ['Accessibility','Monthly',                         'QA + a11y expert if available'],
        ['Performance', 'Monthly + before marketing pushes','QA with Lighthouse + real-device'],
    ],
    col_widths=[3.5, 5.5, 5.5],
)

h2('4.2 Priority classification')

table(
    ['Level', 'Definition', 'SLA'],
    [
        ['P0', 'Blocker — prevents core sign-up/onboarding/sales; data loss risk',    'Fix same day'],
        ['P1', 'High — feature broken; user-visible; affects >20% of users',           'Fix within 3 days'],
        ['P2', 'Medium — feature partially broken; workaround exists',                 'Fix within 2 weeks'],
        ['P3', 'Low — cosmetic; edge case; nice-to-have',                              'Fix when convenient'],
    ],
    col_widths=[1.5, 9.0, 3.5],
)

h2('4.3 Bug report format')

para('When logging a bug, include:', bold=True)
bullet('Title — one sentence, actionable ("Cash-up variance shows negative when it should be zero")')
bullet('Environment — Production / Preview / Local; browser + version; device')
bullet('Steps to reproduce — numbered, with URLs and inputs')
bullet('Expected result — what should happen')
bullet('Actual result — what did happen')
bullet('Screenshot / screen recording — always')
bullet('Console errors — copy from DevTools if any')
bullet('Sentry issue link — if the error is captured')
bullet('Priority — P0/P1/P2/P3 with justification')
bullet('Affected users — estimate (all / role X / plan Y / etc.)')

h2('4.4 Bug tracking')

para(
    'Log bugs in a shared spreadsheet or issue tracker. Include a "Fixed in" '
    'column referencing the Git commit SHA once resolved. Retest and mark '
    'closed only after verification in the release build.',
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. APPENDIX
# ═══════════════════════════════════════════════════════════════════════

h1('5. Appendix')

h2('5.1 Test data seeding')

para(
    'For controlled testing, use the starter-pack products from onboarding. '
    'Avoid adding excessive test data to production. If needed, ask the '
    'founder to seed a dedicated QA account and share credentials.',
)

h2('5.2 Reference URLs')

table(
    ['Purpose', 'URL'],
    [
        ['Production',      'https://www.stokiapp.com'],
        ['Admin console',   'https://www.stokiapp.com/admin (allowlisted only)'],
        ['Sentry',          'https://siti-group-7z.sentry.io/issues/'],
        ['Vercel',          'https://vercel.com/sitigroup-projects/stoki-app'],
        ['Supabase',        'https://supabase.com/dashboard'],
        ['Resend',          'https://resend.com/emails'],
        ['Zoho Mail',       'https://mail.zoho.com'],
        ['Repository',      'https://github.com/Stikili/stoki'],
    ],
    col_widths=[4.0, 11.0],
)

h2('5.3 Contact for issues')

bullet('Product issues / blockers  →  founder (WhatsApp + support@stokiapp.com)')
bullet('Server / infrastructure    →  founder')
bullet('Payment integration Qs     →  founder (Ozow integration TBD)')

h2('5.4 Version history')

table(
    ['Version', 'Date', 'Author', 'Notes'],
    [
        ['1.0', '2026-07-14', 'Founder (via Claude)', 'Initial QA test plan covering all shipped features'],
    ],
    col_widths=[2.5, 3.0, 4.0, 5.5],
)

# ── Save ──────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
print(f'Sections: 5 (About, Environments, Scope, Methodology, Appendix)')
print(f'Test-case checkboxes: ~200 individual test items across 29 feature areas')
